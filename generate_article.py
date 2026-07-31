
import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets the inner padding (margins) of a table cell in dxa (1/20 of a point)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Sets custom borders on a cell. Colors are hex strings (e.g., 'CCCCCC')."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, color in borders.items():
        if color:
            node = OxmlElement(f'w:{border_name}')
            node.set(qn('w:val'), 'single')
            node.set(qn('w:sz'), '12')  # border width (12 = 1.5 pt)
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), color)
            tcBorders.append(node)
        else:
            node = OxmlElement(f'w:{border_name}')
            node.set(qn('w:val'), 'none')
            tcBorders.append(node)
    tcPr.append(tcBorders)

def add_code_block(doc, code_text):
    """Adds a beautifully styled code block (single-cell table with background and borders)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.0)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, 'F4F4F6')  # Soft gray background
    set_cell_margins(cell, top=144, bottom=144, left=216, right=216)
    set_cell_borders(cell, top='E0E0E0', bottom='E0E0E0', left='E0E0E0', right='E0E0E0')
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.0)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)

def add_callout(doc, text):
    """Adds a callout box (single-cell table with a thick left border)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.0)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, 'F0F4F8')  # Very light blue-gray background
    set_cell_margins(cell, top=144, bottom=144, left=216, right=216)
    set_cell_borders(cell, left='2C5E8A', top=None, bottom=None, right=None)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x4E, 0x68)

def format_run_font(run, font_name='Calibri', size_pt=11, bold=False, italic=False, color_rgb=None):
    """Utility to format run fonts easily."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def build_document():
    doc = Document()
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # --- TITLE ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(6)
    title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run("Decentralized Password Management: Building a Web3-Native Vault with AES-256-GCM, IPFS, and BNB Chain")
    format_run_font(run, 'Arial', 22, bold=True, color_rgb=RGBColor(0x11, 0x18, 0x27))
    
    # --- SUBTITLE ---
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(24)
    run = sub_p.add_run("How to combine client-side AES-256-GCM encryption, Web Crypto API deterministic key derivation, IPFS storage, BNB Chain smart contracts, and GraphQL subgraph indexing for 100% user-owned password security.")
    format_run_font(run, 'Calibri', 13.5, italic=True, color_rgb=RGBColor(0x4B, 0x55, 0x63))
    
    # --- INTRO ---
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("In an era dominated by digital accounts, the security of our credentials has never been more critical. Yet, standard solutions continue to rely on centralized frameworks. Traditional password managers—while convenient—store credentials on centralized cloud databases or rely on master keys that present a single point of failure. If the database is breached, or the service provider's servers go down, users face catastrophic security risks or total loss of access.")
    format_run_font(run, 'Calibri', 11)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("Web3 introduces a new paradigm: self-custody and cryptographic identity. What if we could build a password manager where passwords are encrypted client-side using industry-standard ")
    format_run_font(run, 'Calibri', 11)
    run = p.add_run("AES-256-GCM")
    format_run_font(run, 'Calibri', 11, bold=True)
    run = p.add_run(", stored on a decentralized filesystem (IPFS), registered on a secure blockchain network (BNB Chain), and decrypted ")
    format_run_font(run, 'Calibri', 11)
    run = p.add_run("only")
    format_run_font(run, 'Calibri', 11, italic=True)
    run = p.add_run(" by deriving keys client-side directly from the user's Web3 wallet?")
    format_run_font(run, 'Calibri', 11)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("This is the foundation of ")
    format_run_font(run, 'Calibri', 11)
    run = p.add_run("SecureVault")
    format_run_font(run, 'Calibri', 11, bold=True)
    run = p.add_run(". By upgrading from external key-node networks to pure client-side Web Crypto API encryption, SecureVault guarantees zero third-party trust, zero master password friction, and total user sovereignty over digital credentials.")
    format_run_font(run, 'Calibri', 11)
    
    # --- HEADING 1: Tech Stack ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("The Technical Architecture Stack")
    format_run_font(run, 'Arial', 16, bold=True, color_rgb=RGBColor(0x1F, 0x29, 0x37))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("To eliminate centralized vulnerability vectors and eliminate dependency on external key management services, SecureVault leverages a production-grade Web3 stack:")
    format_run_font(run, 'Calibri', 11)
    
    # Tech Stack Table
    tech_table = doc.add_table(rows=7, cols=3)
    tech_table.style = 'Light Shading Accent 1'
    hdr_cells = tech_table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology Used'
    hdr_cells[2].text = 'Primary Role & Purpose'
    
    for cell in hdr_cells:
        set_cell_background(cell, '1F2937')
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        for p_cell in cell.paragraphs:
            for run_cell in p_cell.runs:
                format_run_font(run_cell, 'Arial', 10, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
                
    tech_data = [
        ("Frontend UI", "Next.js 14 & Ant Design (Antd)", "Provides a responsive dashboard for password management, generation, and clipboard integration."),
        ("Web3 Client", "Ethers.js, Wagmi & RainbowKit", "Manages wallet connection, network switching to BSC Testnet, and transaction signing."),
        ("Key Derivation", "MetaMask Signature + SHA-256", "Derives a 256-bit AES CryptoKey deterministically from the user's wallet signature."),
        ("Symmetric Encryption", "Web Crypto API (AES-256-GCM)", "Performs hardware-accelerated client-side encryption with 96-bit random IVs in the browser."),
        ("Decentralized Storage", "IPFS (via Pinata)", "Stores encrypted JSON payloads off-chain to keep gas costs minimal."),
        ("On-Chain Registry", "Solidity (Hardhat on BNB Chain)", "KeyManager smart contract logs IPFS CIDs tied immutably to user wallet addresses.")
    ]
    
    for idx, (comp, tech, role) in enumerate(tech_data, start=1):
        row_cells = tech_table.rows[idx].cells
        row_cells[0].text = comp
        row_cells[1].text = tech
        row_cells[2].text = role
        bg_color = 'F9FAFB' if idx % 2 == 1 else 'FFFFFF'
        for cell in row_cells:
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            set_cell_borders(cell, top='E5E7EB', bottom='E5E7EB', left='E5E7EB', right='E5E7EB')
            for p_cell in cell.paragraphs:
                for run_cell in p_cell.runs:
                    format_run_font(run_cell, 'Calibri', 10)
                    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # --- HEADING 1: Workflow Architecture ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Workflow Architecture: The End-to-End Data Pipeline")
    format_run_font(run, 'Arial', 16, bold=True, color_rgb=RGBColor(0x1F, 0x29, 0x37))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("Plaintext passwords never touch a network socket. The workflow below illustrates how key derivation, AES encryption, IPFS pinning, and blockchain registration coordinate seamlessly:")
    format_run_font(run, 'Calibri', 11)
    
    # Workflow Diagram
    flow_img_path = os.path.join(base_dir, "resources", "lit-pm-flow.png.png")
    if os.path.exists(flow_img_path):
        img_p = doc.add_paragraph()
        img_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(6)
        img_p.paragraph_format.space_after = Pt(4)
        img_run = img_p.add_run()
        img_run.add_picture(flow_img_path, width=Inches(5.5))
        
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(14)
        cap_run = cap_p.add_run("Figure 1: End-to-end cryptographic pipeline and smart contract synchronization in SecureVault.")
        format_run_font(cap_run, 'Calibri', 9, italic=True, color_rgb=RGBColor(0x6B, 0x72, 0x80))
    else:
        add_callout(doc, "[Workflow Diagram: resources/lit-pm-flow.png.png]")
        
    steps = [
        ("1. Wallet Authentication", "The user connects their Web3 wallet (MetaMask) to the BSC Testnet network."),
        ("2. Deterministic Key Derivation", "The user signs a fixed, human-readable key-derivation message. The signature is digested via SHA-256 into a 32-byte AES-256-GCM CryptoKey using `crypto.subtle.importKey`."),
        ("3. Client-Side AES-256-GCM Encryption", "The browser generates a 96-bit random Initialization Vector (IV) and encrypts the credential payload using the derived CryptoKey."),
        ("4. IPFS Archival", "The base64 ciphertext and IV are formatted into a JSON payload `{ ciphertext, iv, version: 'aes-gcm-v1' }` and pinned to IPFS via Pinata, returning a unique IPFS CID."),
        ("5. On-Chain Anchoring", "The client calls `addKey(ipfsHash)` or `updateKey(id, ipfsHash)` on the `KeyManager.sol` smart contract on BNB Chain Testnet."),
        ("6. Graph Subgraph Indexing", "The Graph indexer listens for `KeyAdded` / `KeyUpdated` events, fetches the JSON file via `ipfs.cat`, and caches the payload fields for sub-second GraphQL retrieval.")
    ]
    
    for title, desc in steps:
        sp = doc.add_paragraph(style='List Bullet')
        sp.paragraph_format.space_after = Pt(4)
        sp.paragraph_format.line_spacing = 1.15
        run_title = sp.add_run(f"{title}: ")
        format_run_font(run_title, 'Calibri', 11, bold=True)
        run_desc = sp.add_run(desc)
        format_run_font(run_desc, 'Calibri', 11)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # --- HEADING 1: Deep Dive 1: Key Derivation ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Deep Dive 1: Deterministic Wallet-Based Key Derivation")
    format_run_font(run, 'Arial', 16, bold=True, color_rgb=RGBColor(0x1F, 0x29, 0x37))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("A major bottleneck in Web3 encryption is secret key management. Traditional approaches require users to remember master passwords or rely on external key-share node networks. SecureVault solves this by deriving the encryption key directly from the user's Web3 wallet signature using the Web Crypto API (`client/lib/crypto.js`):")
    format_run_font(run, 'Calibri', 11)
    
    key_code = """// client/lib/crypto.js - Deterministic Key Derivation
const KEY_DERIVATION_MESSAGE = (address) =>
  [
    "SecureVault — Encryption Key Derivation",
    "",
    `Wallet: ${address.toLowerCase()}`,
    "",
    "Signing this message generates your AES-256-GCM encryption key.",
    "It is deterministic: signing again with the same wallet always",
    "produces the same key, so you can always decrypt your passwords.",
    "",
    "⚠️  Never share this signature with anyone.",
  ].join("\\n");

export const deriveEncryptionKey = async (signer, address) => {
  const message = KEY_DERIVATION_MESSAGE(address);
  const signature = await signer.signMessage(message);

  const encoded = new TextEncoder().encode(signature);
  const digest = await crypto.subtle.digest("SHA-256", encoded);

  return crypto.subtle.importKey("raw", digest, { name: "AES-GCM" }, false, [
    "encrypt",
    "decrypt",
  ]);
};"""
    add_code_block(doc, key_code)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("Why is this approach cryptographically sound?")
    format_run_font(run, 'Calibri', 11, bold=True)
    
    reasons = [
        ("Deterministic & Stateless", "Signing the exact same message with the private key produces the exact same signature every time. This means no keys are ever saved to disk or transmitted over the wire."),
        ("Sole Control", "Only the holder of the wallet's private key can produce the required signature, guaranteeing that only the wallet owner can generate the matching AES key."),
        ("No External Node Dependencies", "Unlike threshold networks (e.g. Lit Protocol), this approach runs 100% locally inside the browser with zero network latency and no external server trust requirements.")
    ]
    for r_title, r_desc in reasons:
        sp = doc.add_paragraph(style='List Bullet')
        sp.paragraph_format.space_after = Pt(4)
        sp.paragraph_format.line_spacing = 1.15
        run_t = sp.add_run(f"{r_title}: ")
        format_run_font(run_t, 'Calibri', 11, bold=True)
        run_d = sp.add_run(r_desc)
        format_run_font(run_d, 'Calibri', 11)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # --- HEADING 1: Deep Dive 2: AES-256-GCM ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Deep Dive 2: Native Web Crypto API AES-256-GCM Encryption")
    format_run_font(run, 'Arial', 16, bold=True, color_rgb=RGBColor(0x1F, 0x29, 0x37))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("Galois/Counter Mode (GCM) provides both confidentiality and authenticated encryption (AEAD). Each encryption call generates a fresh 96-bit Initialization Vector (IV), ensuring identical plaintexts yield entirely distinct ciphertexts:")
    format_run_font(run, 'Calibri', 11)
    
    aes_code = """// client/lib/crypto.js - AES-256-GCM Encrypt & Decrypt
export const encryptString = async (cryptoKey, plaintext) => {
  const iv = crypto.getRandomValues(new Uint8Array(12)); // 96-bit random IV
  const encoded = new TextEncoder().encode(plaintext);

  const cipherBuffer = await crypto.subtle.encrypt(
    { name: "AES-GCM", iv },
    cryptoKey,
    encoded
  );

  return {
    ciphertext: toBase64(cipherBuffer),
    iv: toBase64(iv),
    version: "aes-gcm-v1",
  };
};

export const decryptString = async (cryptoKey, { ciphertext, iv }) => {
  const decrypted = await crypto.subtle.decrypt(
    { name: "AES-GCM", iv: fromBase64(iv) },
    cryptoKey,
    fromBase64(ciphertext)
  );
  return new TextDecoder().decode(decrypted);
};"""
    add_code_block(doc, aes_code)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("The generated JSON payload (containing `ciphertext`, `iv`, and `version`) is pinned directly to IPFS, ensuring that anyone inspecting the IPFS hash sees only high-entropy, authenticated ciphertext.")
    format_run_font(run, 'Calibri', 11)
    
    # --- HEADING 1: Smart Contract & Indexer ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Deep Dive 3: The Smart Contract & Subgraph Indexer")
    format_run_font(run, 'Arial', 16, bold=True, color_rgb=RGBColor(0x1F, 0x29, 0x37))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("The ")
    format_run_font(run, 'Calibri', 11)
    run = p.add_run("KeyManager.sol")
    format_run_font(run, 'Calibri', 11, bold=True)
    run = p.add_run(" contract deployed on BNB Chain Testnet acts as an immutable ledger mapping wallet addresses to array indices of IPFS CID hashes:")
    format_run_font(run, 'Calibri', 11)
    
    sol_code = """// SPDX-License-Identifier: MIT
pragma solidity 0.8.20;

contract KeyManager {
    struct Key {
        uint id;
        string ipfsHash;
        bool isDeleted;
    }

    event KeyAdded(uint id, string ipfsHash, address indexed owner);
    event KeyUpdated(uint id, string ipfsHash, address indexed owner);
    event KeyDeleted(uint id, address indexed owner);

    mapping(address => Key[]) keys;

    function addKey(string calldata _ipfsHash) public {
        keys[msg.sender].push(Key(keys[msg.sender].length, _ipfsHash, false));
        emit KeyAdded(keys[msg.sender].length - 1, _ipfsHash, msg.sender);
    }

    function updateKey(uint _id, string calldata _ipfsHash) public {
        keys[msg.sender][_id].ipfsHash = _ipfsHash;
        emit KeyUpdated(_id, _ipfsHash, msg.sender);
    }

    function softDeleteKey(uint _id) public {
        keys[msg.sender][_id].isDeleted = true;
        emit KeyDeleted(_id, msg.sender);
    }
}"""
    add_code_block(doc, sol_code)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("To eliminate slow client-side IPFS gateway fetching, The Graph subgraph indexer listens for contract events, executes `ipfs.cat` during indexing, and caches the encrypted JSON payload for sub-second GraphQL retrieval.")
    format_run_font(run, 'Calibri', 11)
    
    # --- HEADING 1: Comparison Table ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Architectural Comparison: Centralized vs Lit Protocol vs AES-256-GCM")
    format_run_font(run, 'Arial', 16, bold=True, color_rgb=RGBColor(0x1F, 0x29, 0x37))
    
    comp_table = doc.add_table(rows=5, cols=4)
    comp_table.style = 'Light Shading Accent 1'
    c_hdr = comp_table.rows[0].cells
    c_hdr[0].text = 'Feature'
    c_hdr[1].text = 'Centralized (1Password/LastPass)'
    c_hdr[2].text = 'Lit Protocol (Threshold Crypto)'
    c_hdr[3].text = 'SecureVault (Native AES-256-GCM)'
    
    for cell in c_hdr:
        set_cell_background(cell, '1F2937')
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        for p_cell in cell.paragraphs:
            for run_cell in p_cell.runs:
                format_run_font(run_cell, 'Arial', 9.5, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
                
    comp_data = [
        ("Key Derivation", "Master Password + PBKDF2", "Threshold Secret Sharing (TSS) Node Network", "MetaMask Signature + SHA-256 (Deterministic)"),
        ("Encryption Location", "Client/Cloud Server", "Client + Lit Node Protocol", "Pure Browser Web Crypto API (`crypto.subtle`)"),
        ("Third-Party Trust", "High (Cloud DB & Servers)", "Medium (Lit Decentralized Nodes)", "Zero (100% Client-Side Self-Custody)"),
        ("Latency & Cost", "Low / Subscription Fee", "Medium / Node Verification Network", "Sub-second / Low Gas on BNB Chain")
    ]
    
    for idx, (f, c1, c2, c3) in enumerate(comp_data, start=1):
        r_cells = comp_table.rows[idx].cells
        r_cells[0].text = f
        r_cells[1].text = c1
        r_cells[2].text = c2
        r_cells[3].text = c3
        bg_color = 'F9FAFB' if idx % 2 == 1 else 'FFFFFF'
        for cell in r_cells:
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            set_cell_borders(cell, top='E5E7EB', bottom='E5E7EB', left='E5E7EB', right='E5E7EB')
            for p_cell in cell.paragraphs:
                for run_cell in p_cell.runs:
                    format_run_font(run_cell, 'Calibri', 9.5)
                    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # --- HEADING 1: Walkthrough ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("User Interface & Visual Walkthrough")
    format_run_font(run, 'Arial', 16, bold=True, color_rgb=RGBColor(0x1F, 0x29, 0x37))
    
    # Step 1
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    run = h2.add_run("1. Dashboard & Web3 Wallet Connection")
    format_run_font(run, 'Arial', 12, bold=True, color_rgb=RGBColor(0x37, 0x41, 0x51))
    
    home_img = os.path.join(base_dir, "resources", "home-screenshot.png.png")
    if os.path.exists(home_img):
        img_p = doc.add_paragraph()
        img_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_p.add_run()
        img_run.add_picture(home_img, width=Inches(4.8))
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(10)
        cap_run = cap_p.add_run("Figure 2: SecureVault Web3 Dashboard requesting wallet connection.")
        format_run_font(cap_run, 'Calibri', 9, italic=True, color_rgb=RGBColor(0x6B, 0x72, 0x80))
        
    # Step 2
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    run = h2.add_run("2. Password Encryption & Generation")
    format_run_font(run, 'Arial', 12, bold=True, color_rgb=RGBColor(0x37, 0x41, 0x51))
    
    save_img = os.path.join(base_dir, "resources", "save_success-screenshot.png.png")
    if os.path.exists(save_img):
        img_p = doc.add_paragraph()
        img_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_p.add_run()
        img_run.add_picture(save_img, width=Inches(4.8))
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(10)
        cap_run = cap_p.add_run("Figure 3: Save Password dialog with automatic password generator.")
        format_run_font(cap_run, 'Calibri', 9, italic=True, color_rgb=RGBColor(0x6B, 0x72, 0x80))
        
    # Step 3
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    run = h2.add_run("3. Masked Credentials Grid & Editing")
    format_run_font(run, 'Arial', 12, bold=True, color_rgb=RGBColor(0x37, 0x41, 0x51))
    
    edit_img = os.path.join(base_dir, "resources", "edit-password.png.png")
    if os.path.exists(edit_img):
        img_p = doc.add_paragraph()
        img_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_p.add_run()
        img_run.add_picture(edit_img, width=Inches(4.8))
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(10)
        cap_run = cap_p.add_run("Figure 4: Masked credentials table with quick clipboard copy and editing.")
        format_run_font(cap_run, 'Calibri', 9, italic=True, color_rgb=RGBColor(0x6B, 0x72, 0x80))
        
    # Step 4
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    run = h2.add_run("4. On-Chain Smart Contract Updates")
    format_run_font(run, 'Arial', 12, bold=True, color_rgb=RGBColor(0x37, 0x41, 0x51))
    
    update_img = os.path.join(base_dir, "resources", "update_tx-screenshot.png.png")
    if os.path.exists(update_img):
        img_p = doc.add_paragraph()
        img_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_p.add_run()
        img_run.add_picture(update_img, width=Inches(4.8))
        cap_p = doc.add_paragraph()
        cap_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(10)
        cap_run = cap_p.add_run("Figure 5: MetaMask transaction confirmation for on-chain registry updates.")
        format_run_font(cap_run, 'Calibri', 9, italic=True, color_rgb=RGBColor(0x6B, 0x72, 0x80))
        
    # --- HEADING 1: Conclusion ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Conclusion & The Future of Web3 Credentials")
    format_run_font(run, 'Arial', 16, bold=True, color_rgb=RGBColor(0x1F, 0x29, 0x37))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("By combining browser-native AES-256-GCM encryption with deterministic wallet key derivation, IPFS storage, BNB Chain registry contracts, and GraphQL subgraph indexing, SecureVault proves that true cryptographic self-custody is not only possible—it is seamless and highly performant.")
    format_run_font(run, 'Calibri', 11)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run("Users no longer need to rely on vulnerable centralized databases or external key-splitting servers. With SecureVault, your wallet is your master key, the blockchain is your registry, and your passwords belong solely to you.")
    format_run_font(run, 'Calibri', 11)
    
    # Save word document
    out_docx = os.path.join(base_dir, "SecureVault_Medium_Article.docx")
    doc.save(out_docx)
    print(f"SUCCESS: Word Document successfully saved to: {out_docx}")
    
    # Convert to PDF
    out_pdf = os.path.join(base_dir, "SecureVault_Medium_Article.pdf")
    try:
        import win32com.client
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc_obj = word.Documents.Open(out_docx)
        # 17 = wdFormatPDF
        doc_obj.SaveAs(out_pdf, FileFormat=17)
        doc_obj.Close()
        word.Quit()
        print(f"SUCCESS: PDF Document successfully saved to: {out_pdf}")
    except Exception as e:
        print(f"WARNING: Win32com PDF conversion fallback: {e}")
        try:
            from docx2pdf import convert
            convert(out_docx, out_pdf)
            print(f"SUCCESS: PDF Document successfully saved to: {out_pdf} (via docx2pdf)")
        except Exception as e2:
            print(f"ERROR: PDF conversion failed: {e2}")

if __name__ == '__main__':
    build_document()
